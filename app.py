import nest_asyncio
nest_asyncio.apply()
from dotenv import load_dotenv
load_dotenv()
import streamlit as st
import os
from docx import Document
from io import BytesIO
import json

# Gemini and LangChain specific imports
from langchain_community.document_loaders import DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from dotenv import load_dotenv

# Load environment variables for API keys
load_dotenv()

# Check for API key
if "GOOGLE_API_KEY" not in os.environ:
    st.error("Please set your GOOGLE_API_KEY environment variable in a .env file.")
    st.stop()

# Define the checklist for mandatory documents
CHECKLIST = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution Templates",
        "Shareholder Resolution Templates",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors",
    ]
}

# Load ADGM documents for RAG. We use st.cache_resource to run this only once.
@st.cache_resource
def setup_rag():
    st.info("Setting up AI Agent using Gemini...")
    try:
        loader = DirectoryLoader('./adgm_docs', glob="**/*.docx")
        documents = loader.load()
    except Exception as e:
        st.error(f"Error loading documents: {e}")
        st.stop()
    
    # Split documents into smaller, manageable chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    texts = text_splitter.split_documents(documents)
    
    # Create embeddings and a vector store for efficient retrieval using Gemini
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    db = Chroma.from_documents(texts, embeddings)
    
    # Create the LLM instance using Gemini
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro")
    
    # Create the RAG chain
    qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=db.as_retriever())
    
    st.success("AI Agent is ready!")
    return qa

# Setup the RAG chain and run it
qa_chain = setup_rag()

st.title("ADGM Corporate Agent ⚖️")
st.write("Upload your legal documents for review.")

# File uploader widget
uploaded_files = st.file_uploader("Choose DOCX files", type="docx", accept_multiple_files=True)

if uploaded_files:
    st.subheader("Processing Documents...")

    # Determine the legal process based on uploaded documents
    document_names = [file.name for file in uploaded_files]
    process_type = "Company Incorporation" # Can be made dynamic with a more advanced LLM call
    
    # Document Checklist Verification
    uploaded_count = len(uploaded_files)
    required_count = len(CHECKLIST[process_type])
    missing_docs = [doc for doc in CHECKLIST[process_type] if not any(doc.lower() in name.lower() for name in document_names)]
    
    st.write(f"Detected Process: **{process_type}**")
    st.write(f"Uploaded **{uploaded_count}** of **{required_count}** required documents.")
    
    if missing_docs:
        st.warning("Missing Mandatory Documents:")
        for doc in missing_docs:
            st.write(f"- {doc}")

    # Process each uploaded document
    all_issues = []
    
    for file in uploaded_files:
        st.write(f"Analyzing '{file.name}'...")
        
        # Read the document content
        doc = Document(file)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Use LLM to detect red flags with RAG
        prompt = f"""
        Analyze the following legal document content for compliance with ADGM regulations. 
        Identify legal red flags such as incorrect jurisdiction, ambiguous language, or missing clauses.
        For each issue found, provide the clause/section, a description of the issue, and a suggestion for correction.
        Cite the relevant ADGM law or rule if possible.

        Document Content:
        {full_text}

        Provide a response in a JSON format.
        """
        
        try:
            # RAG-powered LLM call to get analysis
            analysis = qa_chain.run(prompt)
            
            # This is a critical step: LLMs can hallucinate. We must validate the output.
            # In a real-world scenario, you would parse the JSON and validate the structure.
            # For this project, we'll assume the LLM provides valid JSON.
            issues_found = json.loads(analysis)
            all_issues.append({"document": file.name, "issues": issues_found})
            
            # Highlight and comment the original DOCX file
            reviewed_doc = Document(file)
            for issue in issues_found:
                for para in reviewed_doc.paragraphs:
                    if issue["section"] in para.text:
                        # A simple way to highlight a paragraph
                        for run in para.runs:
                            run.font.highlight_color = 6 # Yellow highlight
                        para.add_comment(issue["issue"] + " Suggestion: " + issue["suggestion"])
            
            # Create a downloadable version of the reviewed document
            buffer = BytesIO()
            reviewed_doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label=f"Download Reviewed '{file.name}'",
                data=buffer.getvalue(),
                file_name=f"reviewed_{file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Error during analysis of '{file.name}': {e}")
    
    # Generate the final structured output report
    final_report = {
        "process": process_type,
        "documents_uploaded": uploaded_count,
        "required_documents": required_count,
        "missing_document": missing_docs if missing_docs else "None",
        "issues_found": all_issues
    }
    
    st.subheader("Final Structured Report")
    st.json(final_report)
    st.success("Analysis complete!")
